import os
from glob import glob
import pandas as pd
import configparser
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import fitz  # from pyMuPdf package
from PIL import Image, ImageChops
import time

import sys
if len(sys.argv) <= 1:
    print('请输入content.ini文件')
    print('你可以复制和修改 /data/users/dqgu/PycharmProjects/nestcmd/tcr_report/content.ini')
    exit()

content_cfg = sys.argv[1]
if not os.path.exists(content_cfg):
    print('输入文件不存在')
    exit()

config = configparser.ConfigParser(interpolation=configparser.ExtendedInterpolation())
config.optionxform = str

config.read(content_cfg, encoding='utf-8')

document = Document()
style = document.styles['Normal']
document.styles['Normal'].font.name = u'微软雅黑'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
# add new character style for reference
document.styles.add_style('reference', WD_STYLE_TYPE.CHARACTER)
obj_font = document.styles['reference'].font
obj_font.size = Pt(9)
obj_font.name = 'Times New Roman'
# add new heading style
document.styles.add_style('heading', WD_STYLE_TYPE.PARAGRAPH)
obj_font = document.styles['heading'].font
obj_font.name = 'Times New Roman'


# set header
section = document.sections[0]

if config['header']['text']:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = config['header']['text']
    paragraph.alignment = 1

# set footer
if config['header']['text']:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = config['footer']['text']
    paragraph.alignment = 1

# set cover
subject = document.add_heading(config['cover']['subject'], 0)
subject.alignment = 1
fig = document.add_picture(config['cover']['logo'], width=Inches(2))
document.paragraphs[-1].alignment = 1
document.add_paragraph()
document.add_paragraph()
document.add_paragraph('课题负责人：{}'.format(config['cover']['customer'])).alignment = 1
document.add_paragraph('客户单位：{}'.format(config['cover']['department'])).alignment = 1


def pdf2png(pdf, zoom_ratio=2, rotate=0):
    doc = fitz.open(pdf)
    for pg in range(doc.pageCount):
        page = doc[pg]
        trans = fitz.Matrix(zoom_ratio, zoom_ratio).preRotate(rotate)
        # create raster image of page (non-transparent)
        pm = page.getPixmap(matrix=trans, alpha=False)
        # write a PNG image of the page
        out = pdf[:-3] + 'png'
        pm.writePNG(out)
        return out


def trim_white_around(img):
    path = img
    img = Image.open(img)
    bg = Image.new(img.mode, img.size, img.getpixel((0, 0)))
    diff = ImageChops.difference(img, bg)
    diff = ImageChops.add(diff, diff, 2.0, -100)
    bbox = diff.getbbox()
    if bbox:
        img = img.crop(bbox)
    img.save(path)
    return path


def add_table(document, data, top=None):
    data = pd.read_csv(data, header=0, index_col=None, sep=None, engine='python')
    data = data.iloc[:top, :].round(3)
    if top is not None:
        data = data.applymap(lambda x: x[:8]+'...' if type(x) == str and len(x)>=15 else x)
        if len(data.columns) > 10:
            cols = data.columns
            data['...'] = '...'
            data = data[list(cols)[:5] + ['...'] + list(cols)[-5:]]

    table = document.add_table(rows=(data.shape[0]+1), cols=data.shape[1], style='Medium Shading 1 Accent 1')
    table.alignment = 1
    table.style.font.size = Pt(8)
    table.autofit = False
    table.columns[0].width = 8*72*914400 * (len(data.columns[0])+3)
    # add the header rows.
    for j in range(data.shape[1]):
        table.cell(0, j).text = data.columns[j]
        table.cell(0, j).paragraphs[0].alignment = 1
    # add the rest of the data frame
    for i in range(0, data.shape[0]):
        for j in range(data.shape[1]):
            table.cell(i+1, j).text = str(data.values[i, j])
            table.cell(i+1, j).paragraphs[0].alignment = 1


# content
chapters = [x for x in config if x not in ['header', 'footer', 'cover', 'DEFAULT']]
graphs = list()
tables = list()
for chapter in chapters:
    content = config[chapter]
    level, heading = chapter.split(' ', 1)
    level = len([x for x in level.split('.') if x != ''])
    if level == 1:
        document.add_page_break()
    title = document.add_heading('', level=level).add_run(heading)
    title.font.name = u'微软雅黑'
    title._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    print(f'Format section of <{chapter}>')
    for i in range(1, 1+sum(x.startswith('p') for x in content)):
        key = 'p'+str(i)
        if key in content:
            if heading == "关键术语":
                p = document.add_paragraph(content[key], 'List Bullet')
            elif heading == '参考文献':
                p = document.add_paragraph()
                p.add_run('[{}] '.format(i) + content[key], style='reference')
            else:
                p = document.add_paragraph(content[key])
                pf = p.paragraph_format
                pf.first_line_indent = Pt(18)
            if key + '_img' in content:
                figure = glob(content[key+'_img'])
                if figure:
                    figure = figure[0]
                    if figure.endswith('.pdf'):
                        figure = pdf2png(figure)
                    figure = trim_white_around(figure)
                    if heading == '分析流程':
                        document.add_picture(figure, width=Inches(5.8))
                    else:
                        document.add_picture(figure, width=Inches(4.5))
                    document.paragraphs[-1].alignment = 1
                    graphs.append(figure)
                    if key + '_img_caption' in content:
                        caption = '图{} '.format(len(graphs)) + content[key + '_img_caption']
                        if '{sample_name}' in caption:
                            sample_name = os.path.basename(figure).split('.', 1)[0]
                            caption = caption.replace('{sample_name}', sample_name)
                        p = document.add_paragraph(caption, style='Caption')
                        p.alignment = 1
                else:
                    print('Warn: Found no such figure {} in chapter {}'.format(content[key+'_img'], chapter))
                    p = document.add_paragraph('由于样本较少, 不能进行该项分析，当前图片为空', style='Caption')
                    p.alignment = 1

            if key+'_table' in content:
                table = glob(content[key+'_table'])
                if table:
                    if key + '_table_title' in content:
                        title = '表{} '.format(len(tables)+1) + content[key+'_table_title']
                        p = document.add_paragraph(title, style='Caption')
                        p.alignment = 1
                    if table:
                        table = table[0]
                        if key+'_table_display' in content:
                            top = int(content[key+'_table_display'])
                        else:
                            top = None
                        add_table(document, table, top=top)
                        tables.append(table)
                    if key+'_table_caption' in content:
                        caption = '[Note] ' + content[key + '_table_caption']
                        p = document.add_paragraph(caption, style='Caption')
                else:
                    print('Warn: Found no such table {} in chapter {}'.format(content[key+'_table'], chapter))
time_stamp = time.strftime("%Y%m%d", time.localtime())
document.save(f'Report.{time_stamp}.docx')

